"""Unit tests for the IC memo docx renderer."""
from __future__ import annotations

import io
import pytest

from docx import Document

from app.verticals.real_estate.underwriting.memo.docx_renderer import render_memo_docx
from app.verticals.real_estate.underwriting.memo.schemas import (
    Citation,
    MemoContext,
    ProseSection,
    Recommendation,
    Risk,
    RisksSection,
    RetrievedChunk,
    SECTION_EXECUTIVE_SUMMARY,
    SECTION_INVESTMENT_THESIS,
    SECTION_TRANSACTION_OVERVIEW,
    SECTION_PROPERTY_DESCRIPTION,
    SECTION_MARKET_OVERVIEW,
    SECTION_SPONSOR,
    SECTION_FINANCIAL_ANALYSIS,
    SECTION_RENT_POSITION,
    SECTION_RISKS,
    SECTION_RECOMMENDATION,
)


@pytest.fixture
def ctx():
    return MemoContext(
        deal_name="Sunny Storage",
        address="123 Test Ave, Dallas TX",
        asset_type="self_storage",
        year_built=2010,
        num_units=400,
        rentable_sqft=50_000,
        total_unit_count=400,
        storage_unit_count=400,
        non_storage_unit_count=None,
        cc_unit_count=100,
        nc_unit_count=300,
        climate_control_pct=0.25,
        purchase_price=5_000_000.0,
        price_per_unit=12_500.0,
        price_per_sqft=100.0,
        cap_rate_at_cost=0.07,
        population_3mi=60_000,
        avg_household_income_3mi=75_000.0,
        storage_sqft_per_capita_3mi=7.5,
        nearby_storage_1mi=2,
        nearby_storage_3mi=6,
        nearby_storage_5mi=11,
        noi_buildup={"gpr": 800_000.0, "vacancy_loss": 80_000.0, "egi": 720_000.0, "opex": 320_000.0, "noi": 400_000.0},
        return_metrics={"irr": 0.18, "cash_on_cash": 0.09, "equity_multiple": 2.1, "dscr_year_one": 1.45, "debt_yield": 0.092, "break_even_occupancy_pct": 0.78},
        max_loan={"max_loan": 3_250_000.0, "max_loan_by_dscr": 3_948_000.0, "max_loan_by_ltv": 3_250_000.0,
                  "max_loan_by_debt_yield": 5_000_000.0, "binding_constraint": "ltv", "delta_vs_current": 0.0,
                  "current_loan": 3_250_000.0},
        financing={"interest_rate_pct": 0.065, "amortization_years": 25, "loan_term_years": 10},
        classification="Pursue",
        warnings=["DSCR slack thin"],
        cover_data={"deal_name": "Sunny Storage", "prepared_by": "Alice Analyst", "firm": "Acme Capital",
                    "date": "2026-05-20", "address": "123 Test Ave, Dallas TX"},
        noi_bridge={"om_year_one_noi": 410_000.0, "om_current_noi": 395_000.0,
                    "modeled_noi": 400_000.0},
        stress_tests=[
            {"label": "Vacancy +500bps", "scenario_key": "vacancy_plus_500bps",
             "irr": 0.155, "cash_on_cash": 0.059, "dscr_year_one": 1.33, "equity_multiple": 2.41},
            {"label": "Rent Growth → 0%", "scenario_key": "rent_growth_zero",
             "irr": 0.038, "cash_on_cash": 0.076, "dscr_year_one": 1.43, "equity_multiple": 1.32},
        ],
        projections=[
            {"year": 1, "noi": 400_000.0, "debt_service": 250_000.0, "cash_flow": 150_000.0},
            {"year": 2, "noi": 415_200.0, "debt_service": 250_000.0, "cash_flow": 165_200.0},
            {"year": 3, "noi": 430_500.0, "debt_service": 250_000.0, "cash_flow": 180_500.0},
        ],
        capital_structure={
            "purchase_price": 5_000_000.0,
            "down_payment": 1_750_000.0,
            "loan_amount": 3_250_000.0,
            "closing_cost": 100_000.0,
            "capex_reserve_initial": 40_000.0,
            "total_equity_invested": 1_890_000.0,
        },
        rent_position_analysis=[
            {"size": "10 x 10", "climate_type": "NC", "subject_current_rent": 110.0,
             "subject_market_rent": 115.0, "comp_average_rent": 120.0,
             "current_vs_comp_ratio": 0.917, "market_vs_comp_ratio": 0.958, "comp_count": 4,
             "bucket": "medium"},
        ],
        unit_mix=[
            {"size": "10 x 10", "climate_type": "NC", "num_units": 100,
             "occupancy_pct": 0.92, "current_rent": 110.0},
            {"size": "10 x 20", "climate_type": "CC", "num_units": 50,
             "occupancy_pct": 0.88, "current_rent": 175.0},
        ],
        source_support=[
            {
                "group": "Property",
                "field_key": "num_units",
                "label": "Underwriting Unit Count",
                "value": "133",
                "source_basis": "Manual override",
                "citations": "Tulsa Storage OM.pdf: p6",
                "confidence": "-",
                "notes": "Manual override; original value 205. Original source text: Total Units: 205",
            },
            {
                "group": "Debt / Exit",
                "field_key": "exit_cap_rate",
                "label": "Exit Cap Rate",
                "value": "8.61%",
                "source_basis": "Model default",
                "citations": "-",
                "confidence": "0%",
                "notes": "Used default exit cap because OM exit cap was unavailable.",
            },
        ],
    )


def _all_table_text(doc):
    """python-docx table cell text doesn't appear in doc.paragraphs — collect it."""
    parts = []
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


def _sections():
    return {
        SECTION_EXECUTIVE_SUMMARY: ProseSection(paragraphs=["Snapshot.", "Strength.", "Risk + Pursue."]),
        SECTION_INVESTMENT_THESIS: ProseSection(paragraphs=["Stable-income acquisition at attractive going-in yield."]),
        SECTION_TRANSACTION_OVERVIEW: ProseSection(paragraphs=["Purchase at $5M."]),
        SECTION_PROPERTY_DESCRIPTION: ProseSection(
            paragraphs=["400 units, 25% CC.", "Climate controlled per [doc-om-1:p5]."],
            citations=[Citation(doc_id="doc-om-1", page=5)],
        ),
        SECTION_MARKET_OVERVIEW: ProseSection(paragraphs=["Pop 60k.", "Submarket per [doc-om-1:p9]."],
                                              citations=[Citation(doc_id="doc-om-1", page=9)]),
        SECTION_SPONSOR: ProseSection(paragraphs=["Sponsor info not provided."]),
        SECTION_FINANCIAL_ANALYSIS: ProseSection(paragraphs=["NOI $400k, DSCR 1.45x."]),
        SECTION_RENT_POSITION: ProseSection(paragraphs=["In-place rents below market."]),
        SECTION_RISKS: RisksSection(risks=[
            Risk(title="DSCR slack thin", severity="medium", source="verdict_warning",
                 mitigant="Reserve cushion exists."),
            Risk(title="Rollover risk", severity="low", source="rollover", mitigant=None),
            Risk(title="Supply growth", severity="medium", source="analyst_note", mitigant=None),
        ]),
        SECTION_RECOMMENDATION: Recommendation(
            classification="Pursue",
            rationale="DSCR clears floor.",
            driving_metrics=["DSCR 1.45x vs 1.25x floor"],
            conditions=["Monitor DSCR quarterly"],
        ),
    }


class TestRenderMemoDocx:
    def test_returns_valid_docx_bytes(self, ctx):
        out = render_memo_docx(ctx, _sections())
        assert isinstance(out, (bytes, bytearray))
        # Round-trip: openable as a Word document.
        doc = Document(io.BytesIO(out))
        assert len(doc.paragraphs) > 10

    def test_cover_data_appears(self, ctx):
        out = render_memo_docx(ctx, _sections())
        doc = Document(io.BytesIO(out))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Sunny Storage" in full_text
        assert "Alice Analyst" in full_text
        assert "Acme Capital" in full_text

    def test_section_headers_present(self, ctx):
        out = render_memo_docx(ctx, _sections())
        doc = Document(io.BytesIO(out))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        for heading in [
            "Executive Summary",
            "Property Description",
            "Market Overview",
            "Risks & Mitigants",
            "Recommendation",
            "Loan Sizing",
            "Source Citations",
        ]:
            assert heading in full_text, f"Missing heading: {heading}"

    def test_prose_paragraphs_inserted_per_section(self, ctx):
        out = render_memo_docx(ctx, _sections())
        doc = Document(io.BytesIO(out))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Sponsor info not provided" in full_text
        assert "NOI $400k, DSCR 1.45x." in full_text

    def test_risks_rendered_as_numbered_with_severity(self, ctx):
        out = render_memo_docx(ctx, _sections())
        doc = Document(io.BytesIO(out))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "DSCR slack thin" in full_text
        assert "medium" in full_text.lower()

    def test_appendix_lists_cited_docs(self, ctx):
        ctx.citation_doc_labels = {"doc-om-1": "Tulsa Storage OM.pdf"}
        out = render_memo_docx(ctx, _sections())
        doc = Document(io.BytesIO(out))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Tulsa Storage OM.pdf" in full_text
        assert "5" in full_text and "9" in full_text

    def test_appendix_renders_key_input_source_support(self, ctx):
        out = render_memo_docx(ctx, _sections())
        doc = Document(io.BytesIO(out))
        para_text = "\n".join(p.text for p in doc.paragraphs)
        cells = _all_table_text(doc)

        assert "Key Input Source Support" in para_text
        assert "Underwriting Unit Count" in cells
        assert "Manual override" in cells
        assert "original value 205" in cells
        assert "Exit Cap Rate" in cells
        assert "Model default" in cells

    def test_no_unresolved_placeholders(self, ctx):
        out = render_memo_docx(ctx, _sections())
        doc = Document(io.BytesIO(out))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "{{" not in full_text
        assert "}}" not in full_text

    def test_failed_section_renders_placeholder(self, ctx):
        secs = _sections()
        secs[SECTION_RISKS] = None
        out = render_memo_docx(ctx, secs)
        doc = Document(io.BytesIO(out))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "draft manually" in full_text.lower() or "unavailable" in full_text.lower()

    # ── New IC tables ────────────────────────────────────────────────────

    def test_unit_mix_table_rendered(self, ctx):
        out = render_memo_docx(ctx, _sections())
        doc = Document(io.BytesIO(out))
        para_text = "\n".join(p.text for p in doc.paragraphs)
        cells = _all_table_text(doc)
        assert "Unit Mix" in para_text
        assert "Type" in cells
        # Mix-row cells appear inside the table
        assert "10 x 10" in cells
        assert "92.00%" in cells  # occupancy_pct=0.92 -> "92.00%"

    def test_unit_mix_table_uses_category_for_non_storage_rows(self, ctx):
        ctx.unit_mix = [
            {"size": "10 x 20", "unit_category": "covered_parking", "climate_type": "CC",
             "num_units": 2, "occupancy_pct": 0.5, "current_rent": 70.0},
        ]
        out = render_memo_docx(ctx, _sections())
        doc = Document(io.BytesIO(out))
        cells = _all_table_text(doc)
        assert "Covered Parking" in cells

    def test_noi_bridge_table_rendered(self, ctx):
        out = render_memo_docx(ctx, _sections())
        doc = Document(io.BytesIO(out))
        para_text = "\n".join(p.text for p in doc.paragraphs)
        cells = _all_table_text(doc)
        assert "NOI Bridge" in para_text
        assert "OM Year-1 NOI" in cells
        assert "$410,000" in cells

    def test_stress_tests_table_rendered(self, ctx):
        out = render_memo_docx(ctx, _sections())
        doc = Document(io.BytesIO(out))
        para_text = "\n".join(p.text for p in doc.paragraphs)
        cells = _all_table_text(doc)
        assert "Stress Tests" in para_text
        assert "Vacancy +500bps" in cells
        assert "Rent Growth" in cells  # label has unicode arrow; substring is safe

    def test_cash_flow_projection_rendered(self, ctx):
        out = render_memo_docx(ctx, _sections())
        doc = Document(io.BytesIO(out))
        para_text = "\n".join(p.text for p in doc.paragraphs)
        cells = _all_table_text(doc)
        # The h2 includes the dynamic year range
        assert "Cash Flow Projection" in para_text
        assert "Y1" in cells
        assert "Y2" in cells
        assert "$415,200" in cells  # Y2 NOI

    def test_capital_stack_rendered(self, ctx):
        out = render_memo_docx(ctx, _sections())
        doc = Document(io.BytesIO(out))
        para_text = "\n".join(p.text for p in doc.paragraphs)
        cells = _all_table_text(doc)
        assert "Capital Stack" in para_text
        assert "Senior Loan" in cells
        assert "$3,250,000" in cells

    def test_capital_stack_renders_om_proposed_debt_comparison(self, ctx):
        ctx.purchase_price = 2_500_000.0
        ctx.capital_structure = {
            "purchase_price": 2_500_000.0,
            "loan_amount": 1_750_000.0,
            "down_payment": 750_000.0,
            "closing_cost": 50_000.0,
            "total_equity_invested": 800_000.0,
        }
        ctx.om_financing_evidence = {
            "proposed_loan_amount": 1_625_000.0,
            "proposed_down_payment_amount": 875_000.0,
            "proposed_down_payment_pct": 0.35,
            "proposed_ltv_pct": 0.65,
            "model_loan_amount": 1_750_000.0,
            "model_ltv_pct": 0.70,
            "ltv_delta_pct": 0.05,
        }

        out = render_memo_docx(ctx, _sections())
        doc = Document(io.BytesIO(out))
        para_text = "\n".join(p.text for p in doc.paragraphs)
        cells = _all_table_text(doc)

        assert "OM proposed financing" in para_text
        assert "OM Proposed Loan" in cells
        assert "$1,625,000" in cells
        assert "OM Proposed LTV" in cells
        assert "65.00%" in cells
        assert "Model LTV Delta" in cells
        assert "+5.00%" in cells

    def test_transaction_overview_labels_mixed_use_unit_pricing(self, ctx):
        ctx.total_unit_count = 205
        ctx.storage_unit_count = 133
        ctx.non_storage_unit_count = 72
        ctx.price_per_unit = 12_195.0

        out = render_memo_docx(ctx, _sections())
        doc = Document(io.BytesIO(out))
        cells = _all_table_text(doc)

        assert "Price / Total Unit-Space" in cells
        assert "Price / Storage Unit" not in cells

    def test_rent_position_grid_rendered(self, ctx):
        out = render_memo_docx(ctx, _sections())
        doc = Document(io.BytesIO(out))
        para_text = "\n".join(p.text for p in doc.paragraphs)
        cells = _all_table_text(doc)
        assert "Rent Position by Size Bucket" in para_text
        assert "bucket-level weighted averages" in para_text
        assert "Current vs Comp" in cells  # header cell
        assert "Bucket / Example Size" in cells
        assert "Medium bucket (e.g. 10 x 10)" in cells
        assert "10 x 10" in cells

    def test_inline_citation_labels_are_friendly(self, ctx):
        ctx.citation_doc_labels = {"doc-om-1": "Offering Memorandum"}
        out = render_memo_docx(ctx, _sections())
        doc = Document(io.BytesIO(out))
        text = "\n".join(p.text for p in doc.paragraphs)

        assert "[Offering Memorandum:p5]" in text
        assert "[doc-om-1:p5]" not in text

    def test_verdict_override_renders_audit_line(self, ctx):
        """When ctx.verdict_override differs from classification_calculator,
        the Recommendation section shows an italic audit line + reason."""
        ctx.verdict_override = "Pursue"
        ctx.classification_calculator = "Below Screen"
        ctx.verdict_override_reason = "Conversion thesis dominates the static read."
        out = render_memo_docx(ctx, _sections())
        doc = Document(io.BytesIO(out))
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "Analyst override" in text
        assert "Below Screen" in text  # calculator's value is surfaced
        assert "Conversion thesis dominates" in text

    def test_custom_conditions_appear_first(self, ctx):
        """ctx.custom_conditions are prepended to LLM-generated conditions and
        de-duplicated."""
        ctx.custom_conditions = ["Phase I environmental ordered", "Title commitment received"]
        out = render_memo_docx(ctx, _sections())
        doc = Document(io.BytesIO(out))
        text = "\n".join(p.text for p in doc.paragraphs)
        # Custom condition must appear before the LLM-generated one ("Monitor DSCR quarterly").
        custom_idx = text.find("Phase I environmental")
        llm_idx = text.find("Monitor DSCR quarterly")
        assert custom_idx > -1
        assert llm_idx > -1
        assert custom_idx < llm_idx

    def test_strategy_and_sourcing_appear_on_cover(self, ctx):
        ctx.strategy_type = "Value-Add"
        ctx.sourcing_type = "Off-market"
        ctx.sourcing_detail = "Repeat seller"
        out = render_memo_docx(ctx, _sections())
        doc = Document(io.BytesIO(out))
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "Value-Add" in text
        assert "Off-market" in text
        assert "Repeat seller" in text

    def test_section_renumbering_includes_investment_thesis(self, ctx):
        out = render_memo_docx(ctx, _sections())
        doc = Document(io.BytesIO(out))
        text = "\n".join(p.text for p in doc.paragraphs)
        # Investment Thesis is §2; downstream sections shift by one.
        assert "2. Investment Thesis" in text
        assert "3. Transaction Overview" in text
        assert "11. Recommendation" in text
        assert "12. Appendix" in text

    def test_tables_skip_when_data_missing(self, ctx):
        """Empty inputs → render functions early-return, no orphan headers."""
        ctx.stress_tests = []
        ctx.projections = []
        ctx.capital_structure = {}
        ctx.rent_position_analysis = []
        ctx.unit_mix = []
        ctx.noi_bridge = {}
        out = render_memo_docx(ctx, _sections())
        doc = Document(io.BytesIO(out))
        para_text = "\n".join(p.text for p in doc.paragraphs)
        for absent in ["Unit Mix", "NOI Bridge", "Stress Tests",
                       "Cash Flow Projection", "Capital Stack",
                       "Rent Position by Size Bucket"]:
            assert absent not in para_text, f"Header '{absent}' rendered with empty data"
