"""Export utilities: render parsed JSON into Markdown, DOCX, and XLSX and produce PDF via markdown conversion.

Design choices:
- Use python-docx for DOCX, pandas + XlsxWriter for XLSX, and markdown->PDF via pandoc if available or weasyprint.
- Keep the API simple: provide functions that accept parsed JSON (dict) and return bytes + filename.
"""
from __future__ import annotations
import io
import json
from typing import Dict, Any, Tuple

import markdown
from docx import Document
import pandas as pd
from datetime import datetime

try:
    import pypandoc
    PANDOC_AVAILABLE = True
except Exception:
    PANDOC_AVAILABLE = False


def json_to_markdown(parsed: Dict[str, Any]) -> str:
    """Convert structured parsed JSON into a readable markdown memo.

    Simple heuristic: top-level keys become headings; lists become bullet lists; dicts are nested.
    """
    def render_value(v, indent=0):
        pad = ""  # no indentation for markdown block
        if isinstance(v, dict):
            parts = []
            for k, val in v.items():
                parts.append(f"**{k}**: \n\n{render_value(val, indent+1)}")
            return "\n\n".join(parts)
        if isinstance(v, list):
            parts = [f"- {render_value(i, indent+1)}" for i in v]
            return "\n".join(parts)
        return str(v)

    parts = [f"# Generated Memo\n"]
    for k, v in parsed.items():
        parts.append(f"## {k}\n")
        parts.append(render_value(v))
    return "\n\n".join(parts)


def markdown_to_pdf_bytes(md: str) -> bytes:
    """Convert markdown string to PDF bytes. Prefer pypandoc if installed, else raise."""
    if PANDOC_AVAILABLE:
        pdf = pypandoc.convert_text(md, "pdf", format="md")
        return pdf
    else:
        raise RuntimeError("Pandoc not available: install pandoc and pypandoc to enable PDF export")


def to_docx_bytes(parsed: Dict[str, Any]) -> Tuple[bytes, str]:
    doc = Document()
    doc.add_heading('Generated Memo', level=1)
    for k, v in parsed.items():
        doc.add_heading(k, level=2)
        if isinstance(v, list):
            for item in v:
                doc.add_paragraph(str(item), style='List Bullet')
        elif isinstance(v, dict):
            for kk, vv in v.items():
                doc.add_paragraph(f"{kk}: {vv}")
        else:
            doc.add_paragraph(str(v))
    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.read(), "memo.docx"


def to_xlsx_bytes(parsed: Dict[str, Any]) -> Tuple[bytes, str]:
    """Convert parsed workflow JSON into multi-sheet XLSX.

    Heuristics (generic + PE memo friendly):
    - company_info -> sheet CompanyInfo (key/value)
    - financials -> FinancialsYears (pivot of *_by_year dicts) + FinancialsOther (scalar fields)
    - key_risks -> KeyRisks (columns: risk, severity, description, mitigation)
    - management_team -> ManagementTeam (name, title, background, linkedin)
    - transaction_details -> TransactionDetails (key/value)
    - Any remaining top-level dicts -> separate sheet with key/value or list expansion.
    - Lists of dicts become tabular sheet automatically.
    - Raw JSON fallback sheet RawJSON for items not handled.
    """
    bio = io.BytesIO()

    def safe_sheet(name: str) -> str:
        # Excel sheet name limit 31 chars
        return (name[:31]) or "Sheet"

    def write_key_value(writer, sheet: str, data: Dict[str, Any]):
        rows = []
        for k, v in data.items():
            if isinstance(v, (dict, list)):
                try:
                    rows.append((k, json.dumps(v, ensure_ascii=False)))
                except Exception:
                    rows.append((k, str(v)))
            else:
                rows.append((k, v))
        df = pd.DataFrame(rows, columns=["field", "value"])
        df.to_excel(writer, sheet_name=safe_sheet(sheet), index=False)

    def build_financial_years(financials: Dict[str, Any]):
        # Collect year keys from known *_by_year dicts
        year_maps = {
            "revenue": financials.get("revenue_by_year", {}),
            "ebitda": financials.get("ebitda_by_year", {}),
            "adjusted_ebitda": financials.get("adjusted_ebitda_by_year", {}),
            "net_income": financials.get("net_income_by_year", {}),
            "gross_margin": financials.get("gross_margin_by_year", {}),
        }
        all_years = set()
        for m in year_maps.values():
            all_years.update(m.keys())
        rows = []
        for y in sorted(all_years):
            row = {"year": y}
            for label, m in year_maps.items():
                row[label] = m.get(y)
            rows.append(row)
        return pd.DataFrame(rows) if rows else pd.DataFrame(columns=["year"] + list(year_maps.keys()))

    with pd.ExcelWriter(bio, engine="xlsxwriter") as writer:
        # Company Info
        ci = parsed.get("company_info") or parsed.get("companyInformation") or parsed.get("company_information")
        if isinstance(ci, dict) and ci:
            write_key_value(writer, "CompanyInfo", ci)

        # Financials
        fin = parsed.get("financials") or parsed.get("financial_performance")
        if isinstance(fin, dict) and fin:
            df_years = build_financial_years(fin)
            if not df_years.empty:
                df_years.to_excel(writer, sheet_name=safe_sheet("FinancialsYears"), index=False)
            # Other scalar fields (exclude *_by_year dicts)
            scalar_fin = {k: v for k, v in fin.items() if not k.endswith("_by_year")}
            if scalar_fin:
                write_key_value(writer, "FinancialsOther", scalar_fin)

        # Key Risks
        risks = parsed.get("key_risks") or parsed.get("risks")
        if isinstance(risks, list) and risks:
            # Normalize items
            norm_rows = []
            for r in risks:
                if isinstance(r, dict):
                    norm_rows.append({
                        "risk": r.get("risk") or r.get("name") or str(r)[:80],
                        "severity": r.get("severity"),
                        "description": r.get("description"),
                        "mitigation": r.get("mitigation")
                    })
                elif isinstance(r, str):
                    norm_rows.append({"risk": r})
            pd.DataFrame(norm_rows).to_excel(writer, sheet_name=safe_sheet("KeyRisks"), index=False)

        # Management Team
        mgmt = parsed.get("management_team") or parsed.get("management")
        if isinstance(mgmt, list) and mgmt:
            norm_mgmt = []
            for m in mgmt:
                if isinstance(m, dict):
                    norm_mgmt.append({
                        "name": m.get("name") or m.get("title") or str(m)[:60],
                        "title": m.get("title"),
                        "background": m.get("background"),
                        "linkedin": m.get("linkedin")
                    })
                elif isinstance(m, str):
                    norm_mgmt.append({"name": m})
            pd.DataFrame(norm_mgmt).to_excel(writer, sheet_name=safe_sheet("ManagementTeam"), index=False)

        # Transaction Details
        td = parsed.get("transaction_details") or parsed.get("transaction")
        if isinstance(td, dict) and td:
            write_key_value(writer, "TransactionDetails", td)

        # Raw Sections (headings -> text preview)
        raw_sections = parsed.get("raw_sections")
        if isinstance(raw_sections, dict) and raw_sections:
            rows = []
            for heading, val in raw_sections.items():
                text_val = None
                if isinstance(val, dict):
                    text_val = val.get("text") or json.dumps(val)[:400]
                else:
                    text_val = str(val)[:400]
                rows.append({"heading": heading, "text": text_val})
            pd.DataFrame(rows).to_excel(writer, sheet_name=safe_sheet("RawSections"), index=False)

        # Generic fallback for remaining top-level keys not yet handled
        handled = {"company_info", "companyInformation", "company_information", "financials", "financial_performance", "key_risks", "risks", "management_team", "management", "transaction_details", "transaction", "raw_sections"}
        for k, v in parsed.items():
            if k in handled:
                continue
            try:
                if isinstance(v, list) and v and isinstance(v[0], dict):
                    pd.DataFrame(v).to_excel(writer, sheet_name=safe_sheet(k), index=False)
                elif isinstance(v, dict):
                    write_key_value(writer, k, v)
                else:
                    # scalar -> append to Summary sheet
                    pd.DataFrame([(k, v)], columns=["field", "value"]).to_excel(writer, sheet_name=safe_sheet("Summary"), index=False)
            except Exception:
                pd.DataFrame([(k, str(v))], columns=["field", "value"]).to_excel(writer, sheet_name=safe_sheet("Summary"), index=False)

    bio.seek(0)
    return bio.read(), f"memo_{datetime.utcnow().strftime('%Y%m%d_%H%M%S')}.xlsx"


def export_bytes(parsed: Dict[str, Any], fmt: str = 'pdf') -> Tuple[bytes, str, str]:
    """Return (bytes, filename, content_type) for given format: 'pdf'|'docx'|'md'|'xlsx'"""
    md = json_to_markdown(parsed)
    if fmt == 'md':
        return md.encode('utf-8'), 'memo.md', 'text/markdown'
    if fmt == 'docx':
        b, name = to_docx_bytes(parsed)
        return b, name, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    if fmt == 'pdf':
        pdf_bytes = markdown_to_pdf_bytes(md)
        return pdf_bytes, 'memo.pdf', 'application/pdf'
    if fmt == 'xlsx':
        b, name = to_xlsx_bytes(parsed)
        return b, name, 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
    raise ValueError('unsupported format')


__all__ = [
    'json_to_markdown',
    'export_bytes',
]
