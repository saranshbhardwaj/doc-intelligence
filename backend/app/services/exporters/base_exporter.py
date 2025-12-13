"""Base exporter with markdown processing utilities.

Uses markdown + html2docx for robust markdown-to-Word conversion.
"""

from typing import Dict, Any, Tuple
from io import BytesIO
import markdown
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


class BaseExporter:
    """Base class for all exporters with common utilities."""

    def __init__(self):
        self.doc = None

    def create_document(self) -> Document:
        """Create a new Word document."""
        self.doc = Document()
        return self.doc

    def save_to_bytes(self) -> bytes:
        """Save document to bytes."""
        if not self.doc:
            raise ValueError("No document created. Call create_document() first.")

        buffer = BytesIO()
        self.doc.save(buffer)
        buffer.seek(0)
        return buffer.read()

    def add_title(self, text: str, alignment: str = 'center'):
        """Add a title to the document."""
        if not self.doc:
            raise ValueError("No document created")

        para = self.doc.add_heading(text, level=0)
        if alignment == 'center':
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        return para

    def add_heading(self, text: str, level: int = 1):
        """Add a heading."""
        if not self.doc:
            raise ValueError("No document created")
        return self.doc.add_heading(text, level=level)

    def add_paragraph(self, text: str, bold: bool = False, italic: bool = False):
        """Add a paragraph with optional formatting."""
        if not self.doc:
            raise ValueError("No document created")

        para = self.doc.add_paragraph()
        run = para.add_run(text)
        if bold:
            run.bold = True
        if italic:
            run.italic = True
        return para

    def add_markdown_content(self, content: str):
        """
        Convert markdown content to Word document elements.

        Uses markdown library to convert to HTML, then processes HTML.
        For production, consider using html2docx or pypandoc.

        Args:
            content: Markdown-formatted string
        """
        if not content:
            return

        # Parse markdown line by line for better control
        lines = content.split('\n')
        i = 0

        while i < len(lines):
            line = lines[i].strip()

            if not line:
                i += 1
                continue

            # Check for headings (###, ##, #)
            if line.startswith('#'):
                level = len(line) - len(line.lstrip('#'))
                heading_text = line.lstrip('#').strip()
                self.add_heading(heading_text, level=min(level, 3))
                i += 1
                continue

            # Check for bullet lists (-, *)
            if line.startswith(('- ', '* ')):
                bullet_text = line[2:].strip()
                para = self.doc.add_paragraph(style='List Bullet')
                self._add_formatted_text(para, bullet_text)
                i += 1
                continue

            # Check for numbered lists (1., 2., etc.)
            if line[:3].strip() and line[:3].strip()[:-1].isdigit() and line[2:4] == '. ':
                numbered_text = line.split('. ', 1)[1] if '. ' in line else line
                para = self.doc.add_paragraph(style='List Number')
                self._add_formatted_text(para, numbered_text)
                i += 1
                continue

            # Regular paragraph - collect multiline
            para_lines = [line]
            i += 1
            while i < len(lines) and lines[i].strip() and not lines[i].strip().startswith(('#', '-', '*', '1.')):
                para_lines.append(lines[i].strip())
                i += 1

            para_text = ' '.join(para_lines)
            para = self.doc.add_paragraph()
            self._add_formatted_text(para, para_text)

    def _add_formatted_text(self, paragraph, text: str):
        """
        Add text to paragraph with inline markdown formatting (bold, italic).

        Handles:
        - **bold text**
        - *italic text*
        - Plain text
        """
        i = 0
        current_text = ""

        while i < len(text):
            # Check for bold (**text**)
            if i < len(text) - 1 and text[i:i+2] == '**':
                # Add accumulated text first
                if current_text:
                    paragraph.add_run(current_text)
                    current_text = ""

                # Find closing **
                close_idx = text.find('**', i + 2)
                if close_idx != -1:
                    bold_text = text[i+2:close_idx]
                    run = paragraph.add_run(bold_text)
                    run.bold = True
                    i = close_idx + 2
                    continue

            # Check for italic (*text*)
            if text[i] == '*' and (i == 0 or text[i-1] != '*') and (i+1 < len(text) and text[i+1] != '*'):
                # Add accumulated text first
                if current_text:
                    paragraph.add_run(current_text)
                    current_text = ""

                # Find closing *
                close_idx = text.find('*', i + 1)
                if close_idx != -1 and (close_idx+1 >= len(text) or text[close_idx+1] != '*'):
                    italic_text = text[i+1:close_idx]
                    run = paragraph.add_run(italic_text)
                    run.italic = True
                    i = close_idx + 1
                    continue

            # Regular character
            current_text += text[i]
            i += 1

        # Add any remaining text
        if current_text:
            paragraph.add_run(current_text)

    def add_divider(self):
        """Add a horizontal divider."""
        if not self.doc:
            raise ValueError("No document created")
        self.doc.add_paragraph('_' * 50)

    def add_footer_text(self, text: str):
        """Add footer text."""
        if not self.doc:
            raise ValueError("No document created")

        para = self.doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(text)
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(107, 114, 128)  # muted-foreground
        run.italic = True
        return para

    def format_key_value(self, key: str, value: Any) -> str:
        """Format key-value pair for display."""
        if value is None or value == '':
            return f"{key}: -"

        if isinstance(value, bool):
            return f"{key}: {'Yes' if value else 'No'}"

        if isinstance(value, (int, float)):
            return f"{key}: {value:,}"

        if isinstance(value, list):
            return f"{key}: {', '.join(str(v) for v in value)}"

        return f"{key}: {value}"

    def sanitize_filename(self, filename: str, max_length: int = 50) -> str:
        """Sanitize filename for safe download."""
        import re
        # Remove invalid characters
        filename = re.sub(r'[^a-zA-Z0-9._-]', '_', filename)
        # Remove multiple underscores
        filename = re.sub(r'_+', '_', filename)
        # Limit length
        if len(filename) > max_length:
            name, ext = filename.rsplit('.', 1) if '.' in filename else (filename, '')
            name = name[:max_length-len(ext)-1]
            filename = f"{name}.{ext}" if ext else name
        return filename.lower()
