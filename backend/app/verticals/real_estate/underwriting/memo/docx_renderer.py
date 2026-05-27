"""IC memo docx renderer: programmatic Word doc construction with python-docx.

No template file on disk — the document is built from scratch each time. This keeps
the design reviewable in code and removes any binary artifact from the repo.
"""
from __future__ import annotations

import io
import re
from typing import Any, Optional

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt, Inches, RGBColor

from .schemas import (
    MemoContext,
    ProseSection,
    Recommendation,
    RisksSection,
    SECTION_EXECUTIVE_SUMMARY,
    SECTION_INVESTMENT_THESIS,
    SECTION_FINANCIAL_ANALYSIS,
    SECTION_MARKET_OVERVIEW,
    SECTION_PROPERTY_DESCRIPTION,
    SECTION_RECOMMENDATION,
    SECTION_RENT_POSITION,
    SECTION_RISKS,
    SECTION_SPONSOR,
    SECTION_TRANSACTION_OVERVIEW,
)


# ── Helpers ──────────────────────────────────────────────────────────────────

def _h1(doc, text: str) -> None:
    p = doc.add_heading(text, level=1)
    for run in p.runs:
        run.font.size = Pt(16)


def _h2(doc, text: str) -> None:
    doc.add_heading(text, level=2)


def _para(doc, text: str) -> None:
    doc.add_paragraph(text)


def _bold(doc, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True


def _fmt_money(v: Optional[float]) -> str:
    if v is None:
        return "—"
    return f"${v:,.0f}"


def _fmt_pct(v: Optional[float]) -> str:
    if v is None:
        return "—"
    return f"{v * 100:.2f}%"


def _fmt_x(v: Optional[float]) -> str:
    if v is None:
        return "—"
    return f"{v:.2f}x"


def _fmt_int(v: Optional[int]) -> str:
    if v is None:
        return "—"
    return f"{v:,}"


def _add_kv_table(doc, rows: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=len(rows), cols=2)
    for i, (label, value) in enumerate(rows):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[1].text = value


def _display_citation_text(text: str, ctx: MemoContext | None = None) -> str:
    if not ctx or not getattr(ctx, "citation_doc_labels", None):
        return text

    def repl(match: re.Match) -> str:
        doc_id = match.group("doc_id")
        page = match.group("page")
        label = _appendix_source_label(ctx, doc_id)
        return f"[{label}:p{page}]"

    return re.sub(
        r"\[(?P<doc_id>[0-9a-fA-F-]{36}|[A-Za-z0-9_.:-]+):p(?P<page>\d+)\]",
        repl,
        text,
    )


def _render_prose_or_placeholder(doc, section_obj: Optional[Any], ctx: MemoContext | None = None) -> None:
    if section_obj is None:
        p = doc.add_paragraph()
        run = p.add_run("[Section unavailable — please draft manually.]")
        run.italic = True
        run.font.color.rgb = RGBColor(0x99, 0x33, 0x33)
        return
    if isinstance(section_obj, ProseSection):
        for para in section_obj.paragraphs:
            _para(doc, _display_citation_text(para, ctx))


# ── Cover page ───────────────────────────────────────────────────────────────

def _render_cover(doc, ctx: MemoContext) -> None:
    cover = ctx.cover_data or {}
    title = doc.add_paragraph()
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = title.add_run("INVESTMENT COMMITTEE MEMORANDUM")
    run.bold = True
    run.font.size = Pt(20)

    sub = doc.add_paragraph()
    sub.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    sub_run = sub.add_run(cover.get("deal_name", ctx.deal_name))
    sub_run.font.size = Pt(16)

    doc.add_paragraph()  # spacer

    prepared_by = cover.get("prepared_by", "—")
    firm = cover.get("firm", "—")
    address = cover.get("address") or ctx.address or "—"
    date = cover.get("date", "—")

    # Write as paragraphs so they appear in doc.paragraphs (table cells do not).
    cover_rows = [
        ("Property", cover.get("deal_name", ctx.deal_name)),
        ("Address", address),
        ("Prepared by", prepared_by),
        ("Firm", firm),
        ("Date", date),
    ]
    if ctx.strategy_type:
        cover_rows.append(("Strategy", ctx.strategy_type))
    if ctx.sourcing_type:
        sourcing_line = ctx.sourcing_type
        if ctx.sourcing_detail:
            sourcing_line = f"{ctx.sourcing_type} — {ctx.sourcing_detail}"
        cover_rows.append(("Sourcing", sourcing_line))

    for label, value in cover_rows:
        p = doc.add_paragraph()
        p.add_run(f"{label}: ").bold = True
        p.add_run(value or "—")

    doc.add_page_break()


# ── Financial tables ─────────────────────────────────────────────────────────

def _render_noi_buildup(doc, ctx: MemoContext) -> None:
    """Reads calculator's AnnualProjection keys directly. Single vocabulary."""
    n = ctx.noi_buildup or {}
    _add_kv_table(doc, [
        ("Gross Potential Rent", _fmt_money(n.get("gpr"))),
        ("Less: Vacancy / Credit Loss", _fmt_money(n.get("vacancy_loss"))),
        ("Other Income", _fmt_money(n.get("other_income"))),
        ("Effective Gross Income", _fmt_money(n.get("egi"))),
        ("Operating Expenses", _fmt_money(n.get("opex"))),
        ("Net Operating Income", _fmt_money(n.get("noi"))),
    ])


def _render_return_metrics(doc, ctx: MemoContext) -> None:
    """Reads calculator's SelfStorageResult keys directly. Single vocabulary."""
    r = ctx.return_metrics or {}
    _add_kv_table(doc, [
        ("IRR (levered)", _fmt_pct(r.get("irr"))),
        ("Cash-on-Cash", _fmt_pct(r.get("cash_on_cash"))),
        ("Equity Multiple", _fmt_x(r.get("equity_multiple"))),
        ("DSCR (Year 1)", _fmt_x(r.get("dscr_year_one"))),
        ("Debt Yield", _fmt_pct(r.get("debt_yield"))),
        ("Break-even Occupancy", _fmt_pct(r.get("break_even_occupancy_pct"))),
    ])


def _render_loan_sizing(doc, ctx: MemoContext) -> None:
    m = ctx.max_loan or {}
    f = ctx.financing or {}
    _add_kv_table(doc, [
        ("Max Loan by DSCR", _fmt_money(m.get("max_loan_by_dscr"))),
        ("Max Loan by LTV", _fmt_money(m.get("max_loan_by_ltv"))),
        ("Max Loan by Debt Yield", _fmt_money(m.get("max_loan_by_debt_yield"))),
        ("Binding Constraint", str(m.get("binding_constraint") or "—").upper()),
        ("Maximum Supportable Loan", _fmt_money(m.get("max_loan"))),
        ("Current Loan (deal)", _fmt_money(m.get("current_loan"))),
        ("Delta vs. Current", _fmt_money(m.get("delta_vs_current"))),
        ("Proposed Rate", _fmt_pct(f.get("interest_rate_pct"))),
        ("Amortization (years)", _fmt_int(f.get("amortization_years"))),
        ("Loan Term (years)", _fmt_int(f.get("loan_term_years"))),
    ])


# ── Risks ────────────────────────────────────────────────────────────────────

def _render_risks(doc, section_obj: Optional[Any], ctx: MemoContext | None = None) -> None:
    if section_obj is None or not isinstance(section_obj, RisksSection):
        _render_prose_or_placeholder(doc, None, ctx)
        return
    for i, risk in enumerate(section_obj.risks, start=1):
        p = doc.add_paragraph()
        p.add_run(f"{i}. RISK ({risk.severity}): ").bold = True
        p.add_run(_display_citation_text(risk.title, ctx))
        mit = doc.add_paragraph()
        mit.paragraph_format.left_indent = Inches(0.3)
        mit.add_run("MITIGANT: ").bold = True
        mit.add_run(risk.mitigant or "None identified — flag for committee.")


# ── Recommendation ──────────────────────────────────────────────────────────

def _render_recommendation(doc, section_obj: Optional[Any], ctx: MemoContext = None) -> None:
    if section_obj is None or not isinstance(section_obj, Recommendation):
        _render_prose_or_placeholder(doc, None, ctx)
        return

    # Classification line. If analyst overrode the calculator, show the audit trail.
    cls_para = doc.add_paragraph()
    cls_para.add_run("Classification: ").bold = True
    cls_para.add_run(section_obj.classification)
    if ctx and ctx.verdict_override and ctx.classification_calculator and (
        ctx.verdict_override != ctx.classification_calculator
    ):
        override_para = doc.add_paragraph()
        run = override_para.add_run(
            f"Analyst override — calculator classification was "
            f"\"{ctx.classification_calculator}\"."
        )
        run.italic = True
        if ctx.verdict_override_reason:
            reason_para = doc.add_paragraph()
            reason_para.paragraph_format.left_indent = Inches(0.3)
            reason_para.add_run("Override rationale: ").bold = True
            reason_para.add_run(ctx.verdict_override_reason)

    _para(doc, section_obj.rationale)

    if section_obj.driving_metrics:
        _bold(doc, "Driving metrics:")
        for m in section_obj.driving_metrics:
            doc.add_paragraph(m, style="List Bullet")

    # Merge custom_conditions (analyst-supplied, FIRST) with LLM-generated conditions.
    # Both are de-duped (case-insensitive) and capped at 5 total per the prompt.
    merged_conditions: list[str] = []
    if ctx and ctx.custom_conditions:
        merged_conditions.extend(ctx.custom_conditions)
    seen = {c.strip().lower() for c in merged_conditions}
    for c in section_obj.conditions or []:
        key = c.strip().lower()
        if key and key not in seen:
            merged_conditions.append(c)
            seen.add(key)
    merged_conditions = merged_conditions[:5]

    if merged_conditions:
        _bold(doc, "Conditions for proceeding:")
        for c in merged_conditions:
            doc.add_paragraph(c, style="List Bullet")


# ── Additional IC tables ────────────────────────────────────────────────────

def _table_with_header(doc, headers: list[str], rows: list[list[str]]) -> None:
    """Helper: build a styled table with a bold header row and string cells."""
    table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    table.style = "Light List"
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = h
        for r in cell.paragraphs[0].runs:
            r.bold = True
    for i, row in enumerate(rows, start=1):
        for j, val in enumerate(row):
            table.rows[i].cells[j].text = val


def _render_unit_mix(doc, ctx: MemoContext) -> None:
    """Unit-mix table inside Property Description."""
    rows = ctx.unit_mix or []
    if not rows:
        return
    _h2(doc, "Unit Mix")
    body = []
    for u in rows:
        if not isinstance(u, dict):
            continue
        occ = u.get("occupancy_pct") or u.get("occupancy")
        category = str(u.get("unit_category") or "").strip()
        climate = str(u.get("climate_type") or u.get("climate") or "").strip()
        type_label = climate or "—"
        if category and category.lower() != "storage":
            type_label = category.replace("_", " ").title()
        body.append([
            str(u.get("size") or "—"),
            type_label,
            _fmt_int(u.get("num_units")),
            _fmt_pct(occ) if occ is not None else "—",
            _fmt_money(u.get("current_rent")),
        ])
    if not body:
        return
    _table_with_header(doc, ["Size", "Type", "Units", "Occupancy", "Current Rent"], body)


def _render_noi_bridge(doc, ctx: MemoContext) -> None:
    """OM Year-1 NOI vs OM Current NOI vs Model Year-1 NOI."""
    b = ctx.noi_bridge or {}
    if not b:
        return
    om_y1 = b.get("om_year_one_noi") or b.get("om_stated") or b.get("om_year_1")
    om_current = b.get("om_current_noi") or b.get("om_current")
    modeled = b.get("modeled_noi") or b.get("modeled") or b.get("model_year_1")
    if om_y1 is None and om_current is None and modeled is None:
        return
    _h2(doc, "NOI Bridge")
    _add_kv_table(doc, [
        ("OM Year-1 NOI", _fmt_money(om_y1)),
        ("OM Current NOI", _fmt_money(om_current)),
        ("Model Year-1 NOI", _fmt_money(modeled)),
    ])


def _render_stress_tests(doc, ctx: MemoContext) -> None:
    """Stress scenarios as a table: scenario × IRR / CoC / DSCR / EM."""
    tests = ctx.stress_tests or []
    if not tests:
        return
    body = []
    for t in tests:
        if not isinstance(t, dict):
            continue
        body.append([
            str(t.get("label") or t.get("scenario_key") or "—"),
            _fmt_pct(t.get("irr")),
            _fmt_pct(t.get("cash_on_cash")),
            _fmt_x(t.get("dscr_year_one")),
            _fmt_x(t.get("equity_multiple")),
        ])
    if not body:
        return
    _h2(doc, "Stress Tests")
    _table_with_header(
        doc,
        ["Scenario", "IRR", "Cash-on-Cash", "DSCR Y1", "Equity Multiple"],
        body,
    )


def _render_cash_flow_projection(doc, ctx: MemoContext, *, years: int = 5) -> None:
    """First N years of projections: year, NOI, debt service, cash flow."""
    projs = ctx.projections or []
    if not projs:
        return
    rows = projs[:years]
    body = []
    for i, p in enumerate(rows, start=1):
        if not isinstance(p, dict):
            continue
        body.append([
            f"Y{p.get('year', i)}",
            _fmt_money(p.get("noi")),
            _fmt_money(p.get("debt_service")),
            _fmt_money(p.get("cash_flow")),
        ])
    if not body:
        return
    _h2(doc, f"Cash Flow Projection (Years 1-{len(body)})")
    _table_with_header(doc, ["Year", "NOI", "Debt Service", "Cash Flow"], body)


def _render_capital_stack(doc, ctx: MemoContext) -> None:
    """Capital stack / sources & uses at acquisition."""
    cs = ctx.capital_structure or {}
    if not cs:
        return
    _h2(doc, "Capital Stack")
    _add_kv_table(doc, [
        ("Purchase Price", _fmt_money(cs.get("purchase_price"))),
        ("Senior Loan", _fmt_money(cs.get("loan_amount"))),
        ("Equity Down Payment", _fmt_money(cs.get("down_payment"))),
        ("Closing Costs", _fmt_money(cs.get("closing_cost"))),
        ("Capex Reserve (initial)", _fmt_money(cs.get("capex_reserve_initial"))),
        ("Total Equity Invested", _fmt_money(cs.get("total_equity_invested"))),
    ])


def _render_rent_position_grid(doc, ctx: MemoContext) -> None:
    """Per-bucket subject rent vs comp set."""
    rows = ctx.rent_position_analysis or []
    if not rows:
        return
    body = []
    for rp in rows:
        if not isinstance(rp, dict):
            continue
        rp = dict(rp)
        ratio = rp.get("current_vs_comp_ratio")
        bucket = rp.get("bucket")
        size = rp.get("size")
        bucket_label = f"{str(bucket).title()} bucket" if bucket else str(size or "—")
        if bucket and size:
            bucket_label += f" (e.g. {size})"
        rp["size"] = bucket_label
        delta_vs_comp = _fmt_pct(ratio - 1.0) if ratio is not None else "—"
        body.append([
            str(rp.get("size") or "—"),
            str(rp.get("climate_type") or "—"),
            _fmt_money(rp.get("subject_current_rent")),
            _fmt_money(rp.get("subject_market_rent")),
            _fmt_money(rp.get("comp_average_rent")),
            delta_vs_comp,
            _fmt_int(rp.get("comp_count")),
        ])
    if not body:
        return
    _h2(doc, "Rent Position by Size Bucket")
    _para(
        doc,
        "Rows are bucket-level weighted averages, not exact-size rent conclusions. "
        "Use unmatched exact sizes as diligence items before relying on rent-position upside or downside.",
    )
    _table_with_header(
        doc,
        ["Bucket / Example Size", "Climate", "Subject (Current)", "Subject (Market)", "Comp Avg", "Current vs Comp", "Comps"],
        body,
    )


# ── Appendix: citation collection ───────────────────────────────────────────

def _collect_citations(sections: dict[str, Any]) -> dict[str, set[int]]:
    by_doc: dict[str, set[int]] = {}
    for section_obj in sections.values():
        if isinstance(section_obj, ProseSection):
            for cit in section_obj.citations:
                by_doc.setdefault(cit.doc_id, set()).add(cit.page)
        elif isinstance(section_obj, RisksSection):
            for risk in section_obj.risks:
                if risk.citation:
                    by_doc.setdefault(risk.citation.doc_id, set()).add(risk.citation.page)
    return by_doc


_UUID_RE = re.compile(
    r"^[0-9a-f]{8}-[0-9a-f]{4}-[0-9a-f]{4}-[0-9a-f]{4}-[0-9a-f]{12}$",
    re.IGNORECASE,
)


def _appendix_source_label(ctx: MemoContext, doc_id: str) -> str:
    labels = ctx.citation_doc_labels or {}
    label = labels.get(doc_id)
    if label:
        return label
    if doc_id in (ctx.document_ids or []):
        return "Offering Memorandum" if len(ctx.document_ids) == 1 else "Source Document"
    if _UUID_RE.match(str(doc_id or "")):
        return "Source Document"
    return doc_id or "Source Document"


def _render_source_support(doc, ctx: MemoContext) -> None:
    rows = ctx.source_support or []
    if not rows:
        return
    _h2(doc, "Key Input Source Support")
    _para(
        doc,
        "Source support below is taken from the underwriting run's saved field provenance. "
        "Manual and default assumptions may not have a current source page; where available, "
        "manual rows show the original source citation.",
    )
    body = []
    for row in rows:
        if not isinstance(row, dict):
            continue
        body.append([
            str(row.get("group") or "-"),
            str(row.get("label") or row.get("field_key") or "-"),
            str(row.get("value") or "-"),
            str(row.get("source_basis") or "-"),
            str(row.get("citations") or "-"),
            str(row.get("confidence") or "-"),
            str(row.get("notes") or ""),
        ])
    if not body:
        return
    _table_with_header(
        doc,
        ["Group", "Input", "Value Used", "Source Basis", "Citation / Page", "Confidence", "Notes"],
        body,
    )


def _render_appendix(doc, sections: dict[str, Any], ctx: MemoContext) -> None:
    by_doc = _collect_citations(sections)
    if not by_doc:
        _para(doc, "No source citations were used in this memo.")
    else:
        _h2(doc, "Narrative Citations")
        seen_labels: dict[str, int] = {}
        for doc_id, pages in sorted(by_doc.items()):
            base_label = _appendix_source_label(ctx, doc_id)
            count = seen_labels.get(base_label, 0) + 1
            seen_labels[base_label] = count
            label = base_label if count == 1 else f"{base_label} {count}"
            page_list = ", ".join(str(p) for p in sorted(pages))
            _para(doc, f"{label}: pages {page_list}")
    _render_source_support(doc, ctx)


# ── Main entry point ────────────────────────────────────────────────────────

def render_memo_docx(ctx: MemoContext, sections: dict[str, Any]) -> bytes:
    """Render the IC memo to .docx bytes."""
    doc = Document()

    # 0. Cover
    _render_cover(doc, ctx)

    # 1. Executive Summary
    _h1(doc, "1. Executive Summary")
    _render_prose_or_placeholder(doc, sections.get(SECTION_EXECUTIVE_SUMMARY), ctx)

    # 2. Investment Thesis
    _h1(doc, "2. Investment Thesis")
    _render_prose_or_placeholder(doc, sections.get(SECTION_INVESTMENT_THESIS), ctx)

    # 3. Transaction Overview
    _h1(doc, "3. Transaction Overview")
    price_unit_label = "Price / Storage Unit" if ctx.non_storage_unit_count else "Price / Unit"
    _add_kv_table(doc, [
        ("Purchase Price", _fmt_money(ctx.purchase_price)),
        (price_unit_label, _fmt_money(ctx.price_per_unit)),
        ("Price / Rentable Sqft", _fmt_money(ctx.price_per_sqft)),
        ("Going-in Cap Rate", _fmt_pct(ctx.cap_rate_at_cost)),
    ])
    _render_prose_or_placeholder(doc, sections.get(SECTION_TRANSACTION_OVERVIEW), ctx)

    # 4. Property Description
    _h1(doc, "4. Property Description")
    _render_unit_mix(doc, ctx)
    _render_prose_or_placeholder(doc, sections.get(SECTION_PROPERTY_DESCRIPTION), ctx)

    # 5. Market Overview
    _h1(doc, "5. Market Overview")
    _render_prose_or_placeholder(doc, sections.get(SECTION_MARKET_OVERVIEW), ctx)

    # 6. Sponsor / Borrower
    _h1(doc, "6. Sponsor / Borrower")
    _render_prose_or_placeholder(doc, sections.get(SECTION_SPONSOR), ctx)

    # 7. Financial Analysis
    _h1(doc, "7. Financial Analysis")
    _h2(doc, "NOI Buildup")
    _render_noi_buildup(doc, ctx)
    _h2(doc, "Return Metrics")
    _render_return_metrics(doc, ctx)
    _render_noi_bridge(doc, ctx)
    _render_stress_tests(doc, ctx)
    _render_cash_flow_projection(doc, ctx)
    _render_prose_or_placeholder(doc, sections.get(SECTION_FINANCIAL_ANALYSIS), ctx)

    # 8. Rent Position
    _h1(doc, "8. Rent Position")
    _render_prose_or_placeholder(doc, sections.get(SECTION_RENT_POSITION), ctx)
    _render_rent_position_grid(doc, ctx)

    # 9. Loan Sizing & Structure
    _h1(doc, "9. Loan Sizing & Structure")
    _render_loan_sizing(doc, ctx)
    _render_capital_stack(doc, ctx)

    # 10. Risks & Mitigants
    _h1(doc, "10. Risks & Mitigants")
    _render_risks(doc, sections.get(SECTION_RISKS), ctx)

    # 11. Recommendation
    _h1(doc, "11. Recommendation")
    _render_recommendation(doc, sections.get(SECTION_RECOMMENDATION), ctx)

    # 12. Appendix
    _h1(doc, "12. Appendix: Source Citations")
    _render_appendix(doc, sections, ctx)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
